from docx import Document  # importacao o pacote docx para manipular documentos .docx
import time  # para trabalhar com datas
from docx.enum.text import WD_ALIGN_PARAGRAPH  # pacote para formatar os paragrafos
from docx.shared import Pt  # pacote para mudar a fonte e o tamanho delas
from docx.shared import RGBColor  # pacote para mudar a cor da fonte

# -----------------------------------------------
# ---------------------data----------------------
# -----------------------------------------------

def converte_dia(x):
    if x == "0":
        return "domingo"
    if x == "1":
        return "segunda-feira"
    if x == "2":
        return "terça-feira"
    if x == "3":
        return "quarta-feira"
    if x == "4":
        return "quinta-feita"
    if x == "5":
        return "sexta-feita"
    if x == "6":
        return "sábado"


def converte_mes(x):
    if x == "01":
        return "janeiro"
    if x == "02":
        return "fevereiro"
    if x == "03":
        return "março"
    if x == "04":
        return "abril"
    if x == "05":
        return "maio"
    if x == "06":
        return "junho"
    if x == "07":
        return "julho"
    if x == "08":
        return "agosto"
    if x == "09":
        return "setembro"
    if x == "10":
        return "outubro"
    if x == "11":
        return "novembro"
    if x == "12":
        return "dezembro"

def procuracao_aut(nome, cpf, nacionalidade, estado_civil, profissao, endereco, cidade, estado, sexo):
    dia_s = str(time.strftime("%w"))
    dia = str(time.localtime().tm_mday)
    mes = str(time.strftime("%m"))
    ano = str(time.strftime("%Y"))
    domicilio = ""
    if sexo == "F":
        domicilio = "domiciliada"
    else:
        domicilio = "domiciliado"

    procuracao = Document()  # criacao do documento word

    # -----------------------------------------------
    # ---------------------cabeçalho-----------------
    # -----------------------------------------------

    secao_2 = procuracao.sections[0]  # cria o cabeçalho
    cabecalho = secao_2.header
    cabecalho = cabecalho.paragraphs[0]

    cabecalho.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # alinha o texto do cabeçalho à  direita

    logo = cabecalho.add_run()
    # figura = logo.add_picture("logo1.png") #adiciona uma imagem de cabeçalho

    # texto = cabecalho.add_run("	                                                                 Wilson Lopes da Conceição – OAB/PR 21.643\n\
    #	                                                                                                                         Denner Pierro Lourenço – OAB/PR 46.019")


    # fonte_c = texto.font # acessa a fonte do texto do cabeçalho (run)
    # fonte_c.name = "Monotype Corsiva" # fonte do cabeçalho
    # fonte_c.size = Pt(10) # tamanho da fonte

    # cor_c = fonte_c.color # acessa a cor da fonte
    # cor_c.rgb = RGBColor(116,38,51) # muda a cor da fonte usando código RGB

    # -----------------------------------------------
    # ---------------------rodape--------------------
    # -----------------------------------------------

    secao_1 = procuracao.sections[0]  # cria o rodape
    rodape = secao_1.footer
    rodape = rodape.paragraphs[0]

    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER  # centraliza o texto do rodape

    linhas = rodape.add_run("______________________________________________________________________________________\n\
Rua ??????, ?? – Sala ?? – ?º Andar – Fone ?? ????-???? – Londrina – PR\n\
??????@?????.com.br")  # texto do rodape

    fonte_r = linhas.font  # acessa a fonte do texto do rodape (run)
    fonte_r.name = "Monotype Corsiva"  # fonte do rodape
    fonte_r.size = Pt(10)  # tamanho da fonte

    cor_r = fonte_r.color  # acessa a cor da fonte
    cor_r.rgb = RGBColor(153, 51, 0)  # muda a cor da fonte usando código RGB

    # -----------------------------------------------
    # ----------------titulo do documento------------
    # -----------------------------------------------

    cria_titulo = procuracao.add_heading(level=1)  # cria um titulo
    titulo = cria_titulo.add_run('P R O C U R A Ç Ã O')  # adiciona um titulo
    titulo.bold = True  # negrito

    cria_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # centraliza o titulo

    fonte_t = titulo.font  # acessa a fonte do titulo (run)
    fonte_t.name = "Roboto"  # fonte do titulo
    fonte_t.size = Pt(20)  # tamanho da fonte

    cor_t = fonte_t.color  # acessa a cor da fonte (font)
    cor_t.rgb = RGBColor(0, 0, 0)  # muda a cor da fonte usando código RGB

    # ------------------------------------------------
    # ------------- espaços verticais ----------------
    # ------------------------------------------------

    espaco_1 = procuracao.add_paragraph()

    # -----------------------------------------------
    # ---------------- paragrafo 1 ------------------
    # -----------------------------------------------

    paragrafo_1 = procuracao.add_paragraph()  # acrescenta um paragrafo
    outorgante = paragrafo_1.add_run("OUTORGANTE")
    outorgante.bold = True  # negrito
    outorgante.underline = True  # sublinhado
    parte_11 = paragrafo_1.add_run(": ")
    parte_11.bold = True  # negrito
    parte_12 = paragrafo_1.add_run(nome.upper())
    parte_12.bold = True  # negrito
    parte_13 = paragrafo_1.add_run(" (CPF/MF " + cpf + ") " + nacionalidade.lower() + ", " + estado_civil.lower() + ", " + profissao.lower() + ",  residente e " + domicilio + " nesta cidade de "  + cidade + ", \
Estado do " + estado + ", na " + endereco + ".")
    paragrafo_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica os paragrafos

    # ---------- mudança de fonte dos runs ----------------------

    fonte_outorgante = outorgante.font
    fonte_outorgante.name = "Roboto"
    fonte_outorgante.size = Pt(12)

    fonte_11 = parte_11.font
    fonte_11.name = "Roboto"
    fonte_11.size = Pt(12)

    fonte_12 = parte_12.font
    fonte_12.name = "Roboto"
    fonte_12.size = Pt(12)

    fonte_13 = parte_13.font
    fonte_13.name = "Roboto"
    fonte_13.size = Pt(12)

    # --------------------------------------------------------

    # -----------------------------------------------
    # ---------------- paragrafo 2 ------------------
    # -----------------------------------------------

    paragrafo_2 = procuracao.add_paragraph()  # acrescenta um paragrafo
    outorgados = paragrafo_2.add_run("OUTORGADOS")
    outorgados.bold = True  # negrito
    outorgados.underline = True  # sublinhado
    parte_21 = paragrafo_2.add_run(": ")
    parte_21.bold  # negrito
    parte_22 = paragrafo_2.add_run("Nome do advogado".upper())
    parte_22.bold = True  # negrito
    parte_23 = paragrafo_2.add_run(
    ", nacionalidade, estado civil, advogado, inscrito na Ordem dos Advogados do Brasil, Seção ????? (???/?? ?????), ")
    parte_24 = paragrafo_2.add_run("Nome do advogado".upper())
    parte_24.bold = True  # negrito
    parte_25 = paragrafo_2.add_run(
    ", nacionalidade, estado civil, advogado, inscrito na Ordem dos Advogados do Brasil, Seção ????? (???/?? ?????) e ")
    parte_26 = paragrafo_2.add_run("ESCRITÓRIO DE ADVOCACIA")
    parte_26.bold = True  # negrito
    parte_27 = paragrafo_2.add_run(
    " (CNPJ ??.???.???/????-??), todos com escritório profissional na Rua ???????, ??, sala ?? - Fone: (??) ????-????.")
    paragrafo_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica os paragrafos

    # ---------- mudança de fonte dos runs ----------------------

    fonte_outorgados = outorgados.font
    fonte_outorgados.name = "Roboto"
    fonte_outorgados.size = Pt(12)

    fonte_21 = parte_21.font
    fonte_21.name = "Roboto"
    fonte_21.size = Pt(12)

    fonte_22 = parte_22.font
    fonte_22.name = "Roboto"
    fonte_22.size = Pt(12)

    fonte_23 = parte_23.font
    fonte_23.name = "Roboto"
    fonte_23.size = Pt(12)

    fonte_24 = parte_24.font
    fonte_24.name = "Roboto"
    fonte_24.size = Pt(12)

    fonte_25 = parte_25.font
    fonte_25.name = "Roboto"
    fonte_25.size = Pt(12)

    fonte_26 = parte_26.font
    fonte_26.name = "Roboto"
    fonte_26.size = Pt(12)

    fonte_27 = parte_27.font
    fonte_27.name = "Roboto"
    fonte_27.size = Pt(12)

    # ------------------------------------------------


    # -----------------------------------------------
    # ---------------- paragrafo 3 ------------------
    # -----------------------------------------------

    paragrafo_3 = procuracao.add_paragraph()  # acrescenta um paragrafo
    poderes_gerais = paragrafo_3.add_run("PODERES GERAIS")
    poderes_gerais.bold = True  # negrito
    poderes_gerais.underline = True  # sublinhado
    parte_31 = paragrafo_3.add_run(":")
    parte_31.bold = True  # negrito
    parte_32 = paragrafo_3.add_run(" Amplos e ilimitados para o foro em geral, com os da cláusula ")
    ad = paragrafo_3.add_run("ad judicia et extra")
    ad.bold = True  # negrito
    ad.italic = True  # italico
    parte_33 = paragrafo_3.add_run(", podendo representá-lo em juízo ou fora dele, em todas as instâncias judiciais e repartições públicas Federais, \
Estaduais e Municipais, em qualquer ação onde for autor, réu, assistente ou oponente, podendo tudo praticar, requerer, \
assinar, reconvir, concordar, discordar, acordar, ratificar, retificar, acompanhar quaisquer processos e ainda praticar \
todos os demais atos que se fizerem necessários ao integral e fiel cumprimento do presente, podendo SUBSTABELECER, \
no todo ou em parte, com ou sem reserva de poderes.")
    paragrafo_3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica os paragrafos

    # ---------- mudança de fonte dos runs ----------------------

    fonte_poderes_gerais = poderes_gerais.font
    fonte_poderes_gerais.name = "Roboto"
    fonte_poderes_gerais.size = Pt(12)

    fonte_31 = parte_31.font
    fonte_31.name = "Roboto"
    fonte_31.size = Pt(12)

    fonte_32 = parte_32.font
    fonte_32.name = "Roboto"
    fonte_32.size = Pt(12)

    fonte_ad = ad.font
    fonte_ad.name = "Roboto"
    fonte_ad.size = Pt(12)

    fonte_33 = parte_33.font
    fonte_33.name = "Roboto"
    fonte_33.size = Pt(12)

    # --------------------------------------------------------

    # -----------------------------------------------
    # ---------------- paragrafo 4 ------------------
    # -----------------------------------------------

    paragrafo_4 = procuracao.add_paragraph()
    # acrescenta um paragrafo
    poderes_especificos = paragrafo_4.add_run("PODERES ESPECÍFICOS")
    poderes_especificos.bold = True  # negrito
    poderes_especificos.underline = True  # sublinhado
    parte_41 = paragrafo_4.add_run(":")
    parte_41.bold = True  # negrito
    parte_42 = paragrafo_4.add_run(" A presente procuração outorga os poderes especiais para confessar, reconhecer a procedência do pedido, transigir, \
desistir, renunciar ao direito sobre que se funda a ação, firmar compromissos ou acordos, levantar ou receber valores \
através de RPV, PRECATÓRIOS e ALVARÁS, pleitear os benefícios da Gratuidade Judicial, nos termos do Art. 98 do CPC, \
renunciar a valores para fins de determinação de competência e assinar declaração de hipossuficiência econômica, \
tudo em conformidade com a norma do art. 105 do CPC.")
    paragrafo_4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica os paragrafos

# ---------- mudança de fonte dos runs ----------------------

    fonte_poderes_especificos = poderes_especificos.font
    fonte_poderes_especificos.name = "Roboto"
    fonte_poderes_especificos.size = Pt(12)

    fonte_41 = parte_41.font
    fonte_41.name = "Roboto"
    fonte_41.size = Pt(12)

    fonte_42 = parte_42.font
    fonte_42.name = "Roboto"
    fonte_42.size = Pt(12)

    # --------------------------------------------------------

    # ------------------------------------------------
    # ---------------- local e data ------------------
    # ------------------------------------------------


    paragrafo_5 = procuracao.add_paragraph()  # acrescenta um paragrafo
    parte_51 = paragrafo_5.add_run(
    cidade + ", " + converte_dia(dia_s) + ", " + dia + " de " + converte_mes(mes) + " de " + ano + ".")
    paragrafo_5.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Justifica os paragrafos

    # ---------- mudança de fonte dos runs ----------------------

    fonte_51 = parte_51.font
    fonte_51.name = "Roboto"
    fonte_51.size = Pt(12)

    # ----------------------------------------------------------

    # ------------------------------------------------
    # ------------- espaços verticais ----------------
    # ------------------------------------------------

    espaco_2 = procuracao.add_paragraph()
    espaco_3 = procuracao.add_paragraph()

    # ------------------------------------------------
    # ---------------- assinatura --------------------
    # ------------------------------------------------

    paragrafo_6 = procuracao.add_paragraph()
    parte_61 = paragrafo_6.add_run(nome)
    paragrafo_6.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justifica os paragrafos

    # ---------- mudança de fonte dos runs ----------------------

    fonte_61 = parte_61.font
    fonte_61.name = "Roboto"
    fonte_61.size = Pt(12)

    # -----------------------------------------------------------

    # ---------------------------------------------
    # ------ salvamento do arquivo ----------------
    # ---------------------------------------------


    nome_arquivo = "procuração_" + nome + ".docx"

    procuracao.save(nome_arquivo)  # salva o arquivo

    print("Concluído")